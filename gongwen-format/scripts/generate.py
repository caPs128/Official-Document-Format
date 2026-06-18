#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
党政机关公文格式生成器 (GB/T 9704-2012)

将结构化文档内容转换为符合 GB/T 9704-2012 标准的 .docx 文件。

用法:
  1. 修改 build_content(doc) 函数，填入你的文档内容
  2. python generate.py
  3. 打开 output.docx

格式规范:
  - A4纸, 页边距: 上37mm 下35mm 左28mm 右26mm
  - 标题: 2号(22pt) 方正小标宋简体, 居中
  - 正文: 3号(16pt) 仿宋_GB2312, 行距固定值28.9磅, 首行缩进2字符
  - 一级标题: 3号(16pt) 黑体, 加粗
  - 二级标题: 3号(16pt) 楷体_GB2312, 加粗
  - 数字/英文: Times New Roman
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 常量
# ============================================================

SIZE_2HAO = Pt(22)      # 二号
SIZE_3HAO = Pt(16)      # 三号
SIZE_CODE = Pt(10)      # 代码块字号
LINE_SPACING = Pt(28.9) # 固定行距
FIRST_LINE_INDENT = Pt(32)  # 2字符缩进

FONT_TITLE = '方正小标宋简体'
FONT_BODY = '仿宋_GB2312'
FONT_HEI = '黑体'
FONT_KAI = '楷体_GB2312'
FONT_TNR = 'Times New Roman'
FONT_MONO = 'Consolas'

OUTPUT_PATH = 'output.docx'

# ============================================================
# 核心 API
# ============================================================

def set_font(run, east_asian, ascii_font=FONT_TNR, bold=False, size=SIZE_3HAO):
    """
    设置文本 Run 的字体。

    Args:
        run: docx.text.run.Run 对象
        east_asian: 中文字体名称（如 '仿宋_GB2312'）
        ascii_font: 英文/数字字体名称（默认 Times New Roman）
        bold: 是否加粗
        size: 字号（docx.shared.Pt 对象）
    """
    run.font.size = size
    run.bold = bold
    run.font.name = ascii_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:cs'), ascii_font)


def add_paragraph(doc, text, font_east=FONT_BODY, bold=False,
                  indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  before=0, after=0, size=SIZE_3HAO, font_ascii=FONT_TNR):
    """
    添加一个格式完整的正文段落。

    Args:
        doc: Document 对象
        text: 段落文本
        font_east: 中文字体
        bold: 加粗
        indent: 首行缩进（默认 True）
        alignment: 对齐方式
        before: 段前间距（磅）
        after: 段后间距（磅）
        size: 字号
        font_ascii: 英文字体
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING
    pf.alignment = alignment
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.first_line_indent = FIRST_LINE_INDENT

    run = p.add_run(text)
    set_font(run, font_east, font_ascii, bold, size)
    return p


def body(doc, text):
    """正文段落（首行缩进2字符）"""
    return add_paragraph(doc, text)


def body_no_indent(doc, text):
    """正文段落（无缩进，用于列表项）"""
    return add_paragraph(doc, text, indent=False)


def h1_heading(doc, text):
    """一级标题：3号黑体加粗"""
    return add_paragraph(doc, text, font_east=FONT_HEI, bold=True,
                         indent=False, before=12, after=6)


def h2_heading(doc, text):
    """二级标题：3号楷体_GB2312加粗"""
    return add_paragraph(doc, text, font_east=FONT_KAI, bold=True,
                         indent=False, before=8, after=4)


def h3_heading(doc, text):
    """三级标题：3号仿宋_GB2312加粗"""
    return add_paragraph(doc, text, font_east=FONT_BODY, bold=True,
                         indent=False, before=6, after=3)


def doc_title(doc, text):
    """文档主标题：2号方正小标宋简体，居中"""
    return add_paragraph(doc, text, font_east=FONT_TITLE,
                         bold=False, indent=False,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         after=8, size=SIZE_2HAO)


def centered_line(doc, text):
    """居中正文（元数据行）"""
    return add_paragraph(doc, text, indent=False,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2)


def empty_line(doc):
    """空行"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING
    return p


def code_line(doc, text):
    """代码块行（等宽字体，灰底）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    # 灰色背景
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    set_font(run, FONT_BODY, FONT_MONO, size=SIZE_CODE)
    return p


# ============================================================
# 表格 API
# ============================================================

def set_cell(cell, text, font_east=FONT_BODY, bold=False,
             size=SIZE_3HAO, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置表格单元格的文本、字体和内边距"""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_SPACING
    pf.alignment = alignment
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, font_east, FONT_TNR, bold, size)
    # 设置单元格内边距
    tcPr = cell._tc.get_or_add_tcPr()
    mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="40" w:type="dxa"/>'
        f'<w:bottom w:w="40" w:type="dxa"/>'
        f'<w:left w:w="80" w:type="dxa"/>'
        f'<w:right w:w="80" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(mar)


def set_header_cell(cell, text):
    """设置表头单元格（灰色底 + 黑体加粗居中）"""
    set_cell(cell, text, font_east=FONT_HEI, bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="E8EDF5" w:val="clear"/>'
    )
    tcPr.append(shd)


def add_table_borders(table):
    """给表格添加全部边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")} />')
        tbl.insert(0, tblPr)
    # 移除已有边框
    for e in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(e)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows):
    """
    快速创建标准格式表格。

    Args:
        doc: Document 对象
        headers: 表头文字列表，如 ['列1', '列2', '列3']
        rows: 数据行列表，每行是一个文字列表，如 [['a', 'b', 'c'], ['d', 'e', 'f']]

    Returns:
        Table 对象
    """
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        set_header_cell(table.rows[0].cells[i], h)
    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            bold = (c == 0)  # 首列加粗
            set_cell(table.rows[r + 1].cells[c], val, bold=bold)
    add_table_borders(table)
    return table


# ============================================================
# 文档生成入口
# ============================================================

def create_document():
    """创建并配置一个空白的公文文档"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    return doc


def build_content(doc):
    """
    在此函数中构建文档内容。

    可用的辅助函数:
      - doc_title(doc, '标题')          — 居中主标题 (22pt 方正小标宋)
      - centered_line(doc, '副标题')    — 居中文本行
      - h1_heading(doc, '一、章节')     — 一级标题 (黑体)
      - h2_heading(doc, '（一）小节')   — 二级标题 (楷体)
      - h3_heading(doc, '1. 条目')      — 三级标题 (仿宋加粗)
      - body(doc, '正文内容...')        — 正文（首行缩进2字符）
      - body_no_indent(doc, '· 列表')   — 无缩进正文
      - empty_line(doc)                 — 空行
      - code_line(doc, 'mono text')     — 代码块行
      - make_table(doc, headers, rows)  — 标准表格

    示例:
    >>> doc_title(doc, '关于加强人工智能管理的通知')
    >>> empty_line(doc)
    >>> centered_line(doc, '发文机关：某某部委')
    >>> centered_line(doc, '2026年6月18日')
    >>> empty_line(doc)
    >>> h1_heading(doc, '一、总体要求')
    >>> body(doc, '这里是正文内容，首行会自动缩进两个字符...')
    >>> h2_heading(doc, '（一）具体措施')
    >>> body(doc, '进一步展开说明...')
    >>> h3_heading(doc, '1. 技术保障')
    >>> body_no_indent(doc, '· 措施一：加强基础设施建设')
    >>> body_no_indent(doc, '· 措施二：完善数据安全体系')
    >>> make_table(doc,
    ...     ['序号', '任务名称', '责任单位'],
    ...     [['1', '制定标准', '科技司'], ['2', '监督检查', '法规司']])
    """
    # ========== 在此处填写你的文档内容 ==========

    doc_title(doc, '文档标题')
    empty_line(doc)
    centered_line(doc, '副标题或元数据行（可选）')
    empty_line(doc)

    h1_heading(doc, '一、第一章')
    body(doc, '正文内容从这里开始。每段首行自动缩进两个字符。')

    # ========== 内容结束 ==========


# ============================================================
# 主程序
# ============================================================

if __name__ == '__main__':
    doc = create_document()
    build_content(doc)
    doc.save(OUTPUT_PATH)
    print(f'[OK] 公文已生成: {OUTPUT_PATH}')
    print(f'     纸张: A4 (210mm x 297mm)')
    print(f'     页边距: 上37mm 下35mm 左28mm 右26mm')
    print(f'     标题: 22pt 方正小标宋简体')
    print(f'     正文: 16pt 仿宋_GB2312, 行距28.9pt')
    print(f'     首行缩进: 2字符 (32pt)')
